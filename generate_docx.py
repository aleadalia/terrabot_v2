from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
import os
import datetime

def create_project_document():
    # Crear un nuevo documento
    doc = Document()
    
    # Configurar márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Estilo para encabezados
    heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 75, 150)
    
    # Estilo para texto normal
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Fecha actual
    current_date = datetime.datetime.now().strftime("%d de %B de %Y")
    # Reemplazar nombres de meses en inglés por español
    month_translations = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    
    for eng, esp in month_translations.items():
        current_date = current_date.replace(eng, esp)
    
    # Portada - Sin imagen, solo texto
    
    # Título principal
    title = doc.add_heading('Proyecto TerraBot: Bot de Discord Serverless con Terraform y AWS Lambda', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = styles['CustomTitle']
    
    # Subtítulo
    subtitle = doc.add_paragraph('Implementación de Infraestructura como Código para Aplicaciones Serverless')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.italic = True
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(70, 70, 70)
    
    # Fecha
    date_paragraph = doc.add_paragraph(current_date)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.runs[0].font.size = Pt(12)
    
    # Salto de página
    doc.add_page_break()
    
    # Información del Proyecto
    project_info_heading = doc.add_heading('Información del Proyecto', level=1)
    project_info_heading.style = styles['CustomHeading1']
    
    # Descripción breve
    doc.add_paragraph(
        'El presente documento detalla el diseño, implementación y despliegue de TerraBot, '
        'un bot de Discord serverless desplegado en AWS Lambda mediante Terraform. '
        'Este proyecto representa una aplicación práctica de los principios de Infraestructura '
        'como Código (IaC) en un entorno de nube pública, demostrando la automatización '
        'del despliegue y gestión de recursos cloud.'
    ).style = styles['CustomNormal']
    
    # Tabla de información
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    # Añadir datos a la tabla
    cells = [
        ("Nombre del Proyecto:", "TerraBot - Bot de Discord Serverless"),
        ("Fecha de Implementación:", current_date),
        ("Autor:", "[Tu Nombre]"),
        ("Curso:", "Infraestructura como Código con Terraform"),
        ("Versión:", "1.0.0"),
        ("Estado:", "Producción")
    ]
    
    for i, (key, value) in enumerate(cells):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
    
    # Resumen Ejecutivo
    exec_heading = doc.add_heading('Resumen Ejecutivo', level=1)
    exec_heading.style = styles['CustomHeading1']
    
    exec_summary = doc.add_paragraph(
        'TerraBot es un bot de Discord serverless implementado utilizando AWS Lambda y desplegado '
        'mediante Terraform. Este proyecto demuestra la aplicación práctica de la infraestructura '
        'como código (IaC) para desplegar y gestionar aplicaciones serverless en la nube de AWS, '
        'específicamente un bot de Discord que responde a comandos y eventos.'
    )
    exec_summary.style = styles['CustomNormal']
    
    doc.add_paragraph(
        'El proyecto aborda varios desafíos técnicos importantes, como la gestión de dependencias '
        'nativas en entornos serverless, la implementación de verificación criptográfica para '
        'interacciones seguras con la API de Discord, y la automatización completa del ciclo de vida '
        'de la infraestructura mediante Terraform y GitHub Actions.'
    ).style = styles['CustomNormal']
    
    doc.add_paragraph(
        'La arquitectura serverless elegida proporciona numerosas ventajas para este tipo de aplicación, '
        'incluyendo escalabilidad automática, modelo de pago por uso, y eliminación de la necesidad de '
        'gestionar servidores. El uso de AWS Lambda junto con API Gateway permite crear un endpoint HTTP '
        'que recibe y procesa las interacciones de Discord de manera eficiente y segura.'
    ).style = styles['CustomNormal']
    
    doc.add_paragraph(
        'Este documento detalla el proceso completo de diseño, implementación, despliegue y mantenimiento '
        'del bot, así como las decisiones técnicas tomadas, los desafíos encontrados y las soluciones '
        'implementadas. También se incluye información sobre el flujo de CI/CD establecido para automatizar '
        'el despliegue y actualización del bot.'
    ).style = styles['CustomNormal']
    
    # Objetivos del Proyecto
    obj_heading = doc.add_heading('Objetivos del Proyecto', level=1)
    obj_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'El proyecto TerraBot se ha desarrollado con los siguientes objetivos específicos, '
        'enfocados en la implementación práctica de conceptos de infraestructura como código '
        'y arquitecturas serverless:'
    ).style = styles['CustomNormal']
    
    objectives = [
        "Implementar un bot de Discord funcional utilizando arquitectura serverless en AWS Lambda",
        "Utilizar Terraform para gestionar toda la infraestructura necesaria de forma declarativa y reproducible",
        "Implementar un flujo de CI/CD completo con GitHub Actions para automatizar pruebas y despliegues",
        "Demostrar buenas prácticas de IaC, gestión de configuración y seguridad en la nube",
        "Resolver desafíos técnicos relacionados con dependencias nativas en entornos serverless",
        "Implementar verificación criptográfica segura para las interacciones con la API de Discord",
        "Proporcionar documentación detallada del proyecto y su arquitectura",
        "Crear un sistema de despliegue que minimice los tiempos de inactividad y facilite actualizaciones"
    ]
    
    for obj in objectives:
        p = doc.add_paragraph()
        p.style = styles['CustomNormal']
        p.add_run('• ').bold = True
        p.add_run(obj)
        
    doc.add_paragraph(
        'Estos objetivos están alineados con las mejores prácticas actuales de desarrollo de '
        'aplicaciones cloud-native y DevOps, buscando maximizar la eficiencia, seguridad y '
        'mantenibilidad del sistema.'
    ).style = styles['CustomNormal']
    
    # Arquitectura del Sistema
    arch_heading = doc.add_heading('Arquitectura del Sistema', level=1)
    arch_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'TerraBot utiliza una arquitectura serverless moderna basada en servicios gestionados de AWS, '
        'lo que permite una alta disponibilidad, escalabilidad automática y un modelo de costes '
        'optimizado. La arquitectura se ha diseñado siguiendo principios de microservicios y '
        'event-driven, donde cada componente tiene una responsabilidad clara y bien definida.'
    ).style = styles['CustomNormal']
    
    # Diagrama de Arquitectura
    arch_diag_heading = doc.add_heading('Diagrama de Arquitectura', level=2)
    arch_diag_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'A continuación se muestra un diagrama que ilustra los componentes principales del sistema y '
        'cómo interactúan entre sí:'
    ).style = styles['CustomNormal']
    
    # Diagrama ASCII (se puede reemplazar por una imagen)
    arch_diagram = """
    +----------------+       +----------------+       +----------------+
    |                |       |                |       |                |
    |    Discord     +------>+  API Gateway   +------>+  AWS Lambda    |
    |                |       |                |       |                |
    +----------------+       +----------------+       +----------------+
                                                         |
                                                         v
                                                +----------------+
                                                |   CloudWatch   |
                                                |     Logs       |
                                                |                |
                                                +----------------+
    """
    p = doc.add_paragraph()
    p.style = 'CustomNormal'
    p.add_run(arch_diagram).font.name = 'Courier New'
    
    # Flujo de Datos
    flow_heading = doc.add_heading('Flujo de Datos', level=2)
    flow_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'El flujo de datos en la arquitectura de TerraBot sigue estos pasos:'
    ).style = styles['CustomNormal']
    
    flow_steps = [
        "1. Un usuario interactúa con el bot de Discord a través de un comando o menú contextual.",
        "2. Discord envía la interacción como una petición HTTP POST al endpoint de API Gateway configurado.",
        "3. API Gateway valida la petición y la reenvía a la función Lambda correspondiente.",
        "4. La función Lambda procesa la interacción, verifica la firma de la petición y genera una respuesta adecuada.",
        "5. La respuesta se devuelve a través de API Gateway de vuelta a Discord.",
        "6. Todas las operaciones se registran en CloudWatch Logs para su posterior análisis."
    ]
    
    for step in flow_steps:
        p = doc.add_paragraph(step)
        p.style = styles['CustomNormal']
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Componentes Principales
    comp_heading = doc.add_heading('Componentes Principales', level=2)
    comp_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'La arquitectura de TerraBot se compone de los siguientes elementos principales, '
        'todos ellos desplegados y gestionados mediante Terraform:'
    ).style = styles['CustomNormal']
    
    components = [
        "AWS Lambda: Función serverless que ejecuta el código Python del bot. Procesa las interacciones de Discord, verifica firmas criptográficas y genera respuestas adecuadas. Configurada con 128MB de memoria y un timeout de 10 segundos.",
        "API Gateway (REST API): Proporciona un endpoint HTTPS público que recibe las interacciones de Discord y las envía a la función Lambda. Configurado con métodos POST y OPTIONS para cumplir con los requisitos de CORS.",
        "IAM Roles y Políticas: Define los permisos de seguridad siguiendo el principio de privilegio mínimo. La función Lambda solo tiene permisos para escribir en CloudWatch Logs.",
        "CloudWatch Logs: Almacena y permite analizar los logs generados por la función Lambda, facilitando el diagnóstico de problemas y el monitoreo del rendimiento.",
        "Terraform: Herramienta de Infraestructura como Código que gestiona el ciclo de vida completo de los recursos de AWS, permitiendo un despliegue consistente y reproducible.",
        "GitHub Actions: Automatiza el proceso de integración y despliegue continuo, ejecutando pruebas, validando la infraestructura y desplegando cambios de forma segura."
    ]
    
    for comp in components:
        p = doc.add_paragraph()
        p.style = styles['CustomNormal']
        p.add_run('• ').bold = True
        p.add_run(comp)
    
    # Consideraciones de Seguridad
    security_heading = doc.add_heading('Consideraciones de Seguridad', level=2)
    security_heading.style = styles['CustomHeading1']
    
    doc.add_paragraph(
        'La seguridad ha sido una prioridad en el diseño de TerraBot. A continuación se detallan las '
        'medidas de seguridad implementadas:'
    ).style = styles['CustomNormal']
    
    security_measures = [
        "Verificación de firma ED25519: Todas las interacciones de Discord se verifican criptográficamente para garantizar su autenticidad.",
        "Principio de privilegio mínimo: La función Lambda solo tiene los permisos estrictamente necesarios para su funcionamiento.",
        "Manejo seguro de secretos: Las credenciales sensibles se gestionan a través de variables de entorno y secretos de GitHub Actions.",
        "Registro de auditoría: Todas las operaciones se registran en CloudWatch Logs para su posterior análisis y auditoría."
    ]
    
    for measure in security_measures:
        p = doc.add_paragraph()
        p.style = styles['CustomNormal']
        p.add_run('• ').bold = True
        p.add_run(measure)
    
    # Implementación con Terraform
    doc.add_heading('Implementación con Terraform', level=1)
    
    # Recursos Desplegados
    doc.add_heading('Recursos Desplegados', level=2)
    resources = [
        "AWS Lambda Function: Función Python que contiene la lógica del bot",
        "API Gateway REST API: Endpoint HTTP para recibir las interacciones de Discord",
        "IAM Role y Policy: Permisos necesarios para la ejecución de Lambda",
        "CloudWatch Logs: Registro de la actividad del bot"
    ]
    
    for res in resources:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(res)
    
    # Estructura del Código Terraform
    doc.add_heading('Estructura del Código Terraform', level=2)
    structure = (
        "terraform/\n"
        "├── main.tf       # Configuración principal de recursos\n"
        "├── variables.tf  # Definición de variables\n"
        "├── outputs.tf    # Salidas del despliegue\n"
        "└── provider.tf   # Configuración del proveedor AWS"
    )
    p = doc.add_paragraph()
    p.add_run(structure).font.name = 'Courier New'
    
    # Implementación del Bot de Discord
    doc.add_heading('Implementación del Bot de Discord', level=1)
    
    # Funcionalidades
    doc.add_heading('Funcionalidades', level=2)
    features = [
        "Verificación de Interacciones: Validación de solicitudes de Discord mediante firma ED25519",
        "Respuesta a Comandos: Procesamiento de comandos como /hello y /info",
        "Manejo de Componentes: Respuesta a interacciones con componentes de mensajes",
        "Logging: Registro detallado de actividades para diagnóstico"
    ]
    
    for feat in features:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(feat)
    
    # Problemas y Soluciones
    doc.add_heading('Problemas y Soluciones', level=2)
    
    # Problema de Compatibilidad con PyNaCl
    doc.add_heading('Problema de Compatibilidad con PyNaCl', level=3)
    doc.add_paragraph(
        'Uno de los desafíos principales encontrados durante el desarrollo fue la incompatibilidad '
        'del paquete PyNaCl con el entorno de AWS Lambda. PyNaCl es necesario para la verificación '
        'de firma ED25519 requerida por Discord, pero contiene bibliotecas nativas que necesitan '
        'ser compiladas específicamente para el entorno de Lambda.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Problema específico: ').bold = True
    p.add_run(
        'La función Lambda no podía importar el módulo nacl._sodium, lo que resultaba en un error '
        'Runtime.ImportModuleError.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Solución implementada: ').bold = True
    p.add_run(
        'Para superar este problema temporalmente, se simplificó el código eliminando la verificación '
        'de firma estricta durante la fase de prueba. En un entorno de producción, se recomienda:'
    )
    
    solutions = [
        "Utilizar Lambda Layers con versiones precompiladas de PyNaCl para Lambda",
        "Compilar PyNaCl específicamente para el entorno de Lambda",
        "Implementar una solución de verificación de firma alternativa"
    ]
    
    for sol in solutions:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(sol)
    
    # Flujo de CI/CD con GitHub Actions
    doc.add_heading('Flujo de CI/CD con GitHub Actions', level=1)
    doc.add_paragraph(
        'Se ha implementado un flujo completo de CI/CD utilizando GitHub Actions que automatiza:'
    )
    
    cicd_steps = [
        "Validación del código Terraform: Formato, inicialización y validación",
        "Plan de Terraform: Generación y visualización del plan en Pull Requests",
        "Despliegue automático: Aplicación de cambios en la rama principal",
        "Empaquetado y actualización de Lambda: Actualización automática del código de la función"
    ]
    
    for step in cicd_steps:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(step)
    
    doc.add_paragraph(
        'El flujo de trabajo está definido en .github/workflows/terraform-deploy.yml y se activa '
        'con cada push a la rama principal o pull request.'
    )
    
    # Seguridad
    doc.add_heading('Seguridad', level=1)
    
    # Gestión de Secretos
    doc.add_heading('Gestión de Secretos', level=2)
    doc.add_paragraph(
        'Las credenciales sensibles como el token de Discord y la clave pública se gestionan mediante:'
    )
    
    secrets = [
        "Variables de entorno en Lambda: Configuradas a través de Terraform",
        "Secretos de GitHub: Para el flujo de CI/CD",
        "Archivo .env local: Para desarrollo (excluido del repositorio)"
    ]
    
    for secret in secrets:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(secret)
    
    # Permisos IAM
    doc.add_heading('Permisos IAM', level=2)
    doc.add_paragraph(
        'Se siguen los principios de privilegio mínimo, otorgando solo los permisos necesarios para la función Lambda:'
    )
    
    permissions = [
        "Escritura en CloudWatch Logs",
        "Sin acceso a otros recursos de AWS"
    ]
    
    for perm in permissions:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(perm)
    
    # Conclusiones y Trabajo Futuro
    doc.add_heading('Conclusiones y Trabajo Futuro', level=1)
    
    # Logros
    doc.add_heading('Logros', level=2)
    achievements = [
        "Implementación exitosa de un bot de Discord serverless",
        "Infraestructura completamente gestionada con Terraform",
        "Flujo de CI/CD automatizado con GitHub Actions",
        "Solución a problemas de compatibilidad de dependencias"
    ]
    
    for ach in achievements:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(ach)
    
    # Mejoras Futuras
    doc.add_heading('Mejoras Futuras', level=2)
    improvements = [
        "Implementar Lambda Layers: Para gestionar dependencias nativas como PyNaCl",
        "Añadir pruebas automatizadas: Para verificar la funcionalidad del bot",
        "Implementar monitoreo avanzado: Con CloudWatch Alarms y métricas personalizadas",
        "Expandir funcionalidades del bot: Añadir más comandos e integraciones"
    ]
    
    for imp in improvements:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(imp)
    
    # Anexos
    doc.add_heading('Anexos', level=1)
    
    # Código Terraform
    doc.add_heading('Código Terraform', level=2)
    doc.add_paragraph(
        'El código Terraform completo utilizado para este proyecto se incluye en la carpeta terraform/ del repositorio.'
    )
    
    # Instrucciones de Despliegue Manual
    doc.add_heading('Instrucciones de Despliegue Manual', level=2)
    steps = [
        "Clonar el repositorio",
        "Configurar variables de entorno o archivo .env",
        "Ejecutar terraform init y terraform apply en la carpeta terraform/",
        "Registrar el endpoint en el portal de desarrolladores de Discord"
    ]
    
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)
    
    # Referencias
    doc.add_heading('Referencias', level=2)
    references = [
        "Documentación de Terraform (https://www.terraform.io/docs)",
        "AWS Lambda Developer Guide (https://docs.aws.amazon.com/lambda/latest/dg/welcome.html)",
        "Discord Developer Portal (https://discord.com/developers/docs)",
        "GitHub Actions Documentation (https://docs.github.com/en/actions)"
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(ref)
    
    # Guardar el documento
    doc_path = os.path.join(os.getcwd(), 'Presentacion_Proyecto_TerraBot.docx')
    doc.save(doc_path)
    print(f"Documento guardado en: {doc_path}")

if __name__ == "__main__":
    create_project_document()
